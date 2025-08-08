
import sys
import traceback
from flask import Flask, request, send_file, render_template, request, redirect, url_for
import os
import mammoth
from weasyprint import HTML

# 新增依赖
import openpyxl
from xlsx2html import xlsx2html
import io

from PIL import Image

import pdfplumber
import pytesseract
from docx import Document

from pdf2image import convert_from_path
import zipfile

from flask_babel import Babel, _
from babel_config import LANGUAGES

app = Flask(__name__)
app.config['BABEL_DEFAULT_LOCALE'] = 'zh'
app.config['BABEL_TRANSLATION_DIRECTORIES'] = 'translations'
babel = Babel(app)

# 兼容 PyInstaller 打包路径
def resource_path(relative_path):
    if hasattr(sys, '_MEIPASS'):
        base_path = sys._MEIPASS
    else:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

UPLOAD_FOLDER = resource_path('uploads')
CONVERTED_FOLDER = resource_path('converted')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CONVERTED_FOLDER, exist_ok=True)

@app.route('/')
def index():
    lang = request.args.get('lang')
    # 直接渲染页面，避免死循环重定向
    return render_template('index.html', lang=lang or 'zh', languages=LANGUAGES)

@app.route('/convert', methods=['POST'])
def convert_file():
    file = request.files['file']
    if file and file.filename.endswith('.docx'):
        docx_path = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(docx_path)

        # 使用 mammoth 将 docx 转为 HTML
        with open(docx_path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html_content = result.value

        # 使用 weasyprint 将 HTML 转为 PDF
        pdf_filename = file.filename.replace('.docx', '.pdf')
        pdf_path = os.path.join(CONVERTED_FOLDER, pdf_filename)
        HTML(string=html_content).write_pdf(pdf_path)

        return send_file(pdf_path, as_attachment=True)
    else:
        return "请上传 .docx 文件"

# 新增 Excel 转 PDF 路由
@app.route('/convert_excel', methods=['POST'])
def convert_excel():
    file = request.files['file']
    if not (file and file.filename.endswith('.xlsx')):
        return "请上传 .xlsx 文件"

    xlsx_path = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(xlsx_path)

    # 使用 openpyxl 读取 Excel 内容和样式
    wb = openpyxl.load_workbook(xlsx_path)
    ws = wb.active

    html_rows = []
    for row in ws.iter_rows():
        html_cells = []
        for cell in row:
            value = cell.value if cell.value is not None else ""
            style = ""

            # 提取字体样式
            if cell.font:
                if cell.font.bold:
                    style += "font-weight: bold;"
                if cell.font.italic:
                    style += "font-style: italic;"
                if cell.font.color and cell.font.color.rgb:
                    rgb = cell.font.color.rgb
                    if rgb and len(rgb) == 8:  # ARGB 格式
                        style += f"color: #{rgb[2:]};"

            # 提取背景色
            if cell.fill and cell.fill.fgColor and cell.fill.fgColor.rgb:
                bg_rgb = cell.fill.fgColor.rgb
                if bg_rgb and len(bg_rgb) == 8:
                    style += f"background-color: #{bg_rgb[2:]};"

            html_cells.append(f'<td style="{style}">{value}</td>')
        html_rows.append(f"<tr>{''.join(html_cells)}</tr>")

    html_table = f"<table>{''.join(html_rows)}</table>"

    # 构建完整 HTML 页面
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <style>
            body {{
                font-family: "DejaVu Sans", "Arial", sans-serif;
                font-size: 14px;
                padding: 20px;
            }}
            table {{
                border-collapse: collapse;
                width: 100%;
            }}
            th, td {{
                border: 1px solid #ccc;
                padding: 8px;
                text-align: left;
            }}
        </style>
    </head>
    <body>
        <h2>Excel 转 PDF</h2>
        {html_table}
    </body>
    </html>
    """

    # 生成 PDF
    pdf_filename = file.filename.replace('.xlsx', '.pdf')
    pdf_path = os.path.join(CONVERTED_FOLDER, pdf_filename)
    HTML(string=html_content).write_pdf(pdf_path)

    return send_file(pdf_path, as_attachment=True)


# 新增 图片转 PDF 路由
@app.route('/convert_images', methods=['POST'])
def convert_images():
    files = request.files.getlist('files')
    images = []
    for file in files:
        if file and file.filename.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp', '.gif')):
            img = Image.open(file.stream).convert('RGB')
            images.append(img)
    if not images:
        return "请上传图片文件（png/jpg/jpeg/bmp/gif）"
    # 纵向拼接所有图片
    widths, heights = zip(*(img.size for img in images))
    max_width = max(widths)
    total_height = sum(heights)
    merged_img = Image.new('RGB', (max_width, total_height), (255, 255, 255))
    y_offset = 0
    for img in images:
        merged_img.paste(img, (0, y_offset))
        y_offset += img.height
    # 保存为PDF
    pdf_filename = 'images_merged.pdf'
    pdf_path = os.path.join(CONVERTED_FOLDER, pdf_filename)
    merged_img.save(pdf_path, 'PDF')
    return send_file(pdf_path, as_attachment=True)

# 新增 PDF 转 Word/Excel 路由
@app.route('/convert_pdf', methods=['POST'])
def convert_pdf():
    file = request.files['file']
    if not (file and file.filename.endswith('.pdf')):
        return "请上传 PDF 文件"
    pdf_path = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(pdf_path)

    # 选择输出类型
    output_type = request.form.get('output_type', 'docx')  # 'docx' 或 'xlsx'

    # 文本和表格提取
    with pdfplumber.open(pdf_path) as pdf:
        if output_type == 'docx':
            doc = Document()
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    doc.add_paragraph(text)
                # 提取表格
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        doc.add_paragraph('\t'.join([str(cell) for cell in row]))
            docx_filename = file.filename.replace('.pdf', '.docx')
            docx_path = os.path.join(CONVERTED_FOLDER, docx_filename)
            doc.save(docx_path)
            return send_file(docx_path, as_attachment=True)
        elif output_type == 'xlsx':
            import openpyxl
            import re
            wb = openpyxl.Workbook()
            ws = wb.active
            for page in pdf.pages:
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        ws.append(row)
            # 文件名安全处理：只保留中英文、数字、下划线
            safe_name = re.sub(r'[^\w\u4e00-\u9fa5]', '_', os.path.splitext(file.filename)[0])
            xlsx_filename = f"{safe_name}.xlsx"
            xlsx_path = os.path.join(CONVERTED_FOLDER, xlsx_filename)
            wb.save(xlsx_path)
            # 检查文件是否存在
            if not os.path.exists(xlsx_path):
                return f"Excel 保存失败，路径：{xlsx_path}，请检查文件名和磁盘权限。"
            return send_file(xlsx_path, as_attachment=True)

    # 如果没有文本和表格，尝试 OCR
    from pdf2image import convert_from_path
    images = convert_from_path(pdf_path)
    ocr_text = ''
    for img in images:
        ocr_text += pytesseract.image_to_string(img) + '\n'
    doc = Document()
    doc.add_paragraph(ocr_text)
    docx_filename = file.filename.replace('.pdf', '_ocr.docx')
    docx_path = os.path.join(CONVERTED_FOLDER, docx_filename)
    doc.save(docx_path)
    return send_file(docx_path, as_attachment=True)

# 新增 PDF 转图片路由
@app.route('/pdf_to_images', methods=['POST'])
def pdf_to_images():
    file = request.files['file']
    if not (file and file.filename.endswith('.pdf')):
        return "请上传 PDF 文件"
    pdf_path = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(pdf_path)

    images = convert_from_path(pdf_path)
    image_paths = []
    for idx, img in enumerate(images):
        img_filename = f"{os.path.splitext(file.filename)[0]}_page{idx+1}.png"
        img_path = os.path.join(CONVERTED_FOLDER, img_filename)
        img.save(img_path, 'PNG')
        image_paths.append(img_path)

    # 打包为 zip
    zip_filename = f"{os.path.splitext(file.filename)[0]}_images.zip"
    zip_path = os.path.join(CONVERTED_FOLDER, zip_filename)
    with zipfile.ZipFile(zip_path, 'w') as zipf:
        for img_path in image_paths:
            zipf.write(img_path, os.path.basename(img_path))

    return send_file(zip_path, as_attachment=True)

# 新增 PDF 转长图路由
@app.route('/pdf_to_long_image', methods=['POST'])
def pdf_to_long_image():
    import re
    file = request.files['file']
    if not (file and file.filename.endswith('.pdf')):
        return "请上传 PDF 文件"
    pdf_path = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(pdf_path)

    images = convert_from_path(pdf_path)
    widths, heights = zip(*(img.size for img in images))
    total_height = sum(heights)
    max_width = max(widths)
    long_img = Image.new('RGB', (max_width, total_height), (255, 255, 255))
    y_offset = 0
    for img in images:
        long_img.paste(img, (0, y_offset))
        y_offset += img.height
    # 文件名安全处理：只保留中英文、数字、下划线
    safe_name = re.sub(r'[^\w\u4e00-\u9fa5]', '_', os.path.splitext(file.filename)[0])
    long_img_filename = f"{safe_name}_long.png"
    long_img_path = os.path.join(CONVERTED_FOLDER, long_img_filename)
    long_img.save(long_img_path)
    # 检查文件是否存在
    if not os.path.exists(long_img_path):
        return f"长图保存失败，路径：{long_img_path}，请检查文件名和磁盘权限。"
    return send_file(long_img_path, as_attachment=True)

if __name__ == '__main__':
    try:
        port = int(os.environ.get('PORT', 5000))
        app.run(host='0.0.0.0', port=port)
    except Exception as e:
        with open("error.log", "w", encoding="utf-8") as f:
            traceback.print_exc(file=f)
        raise

